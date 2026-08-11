"""
Ricava il template `relazione-dm329.docx` dal documento formattato in Word.

Perché esiste
-------------
La composizione della relazione è chiusa; la formattazione la decide il redattore in Word.
Rigenerare il documento da uno script Python (com'era fino alla Fase 5) distruggerebbe ogni
volta quella formattazione. Ma il documento formattato è un documento *reso*: contiene i
dati di un impianto di esempio, righe di tabella ripetute e un solo ramo di ogni frase
condizionale.

Questo script fa il passaggio inverso: prende il documento formattato e ne ricava il
template, cioè
  - sostituisce i valori con i tag di docxtemplater,
  - collassa le righe ripetute in una sola riga avvolta in un loop,
  - ricrea i rami condizionali che nel documento reso non compaiono,
  - rimette al posto dell'immagine il tag `{%schemaImpianto}`.

Ogni sostituzione dichiara il testo che si aspetta di trovare e **fallisce** se non lo
trova: meglio interrompersi che produrre in silenzio un template mutilo.

I punti d'aggancio si cercano per **testo**, non per posizione. Ancorarli al numero di
paragrafo aveva chiuso il giro: recependo una revisione, l'indirizzo di copertina è passato
da due paragrafi a uno con l'a capo dentro, e da lì in poi ogni conteggio slittava di uno.
Cercando per testo, il documento può guadagnare o perdere paragrafi senza che lo script se
ne accorga, e sono ammesse entrambe le forme dell'indirizzo.

Uso
---
    python scripts/tag-relazione-template.py [sorgente.docx]

Il flusso di lavoro previsto:

  1. npx tsx scripts/generate-relazione-sample.ts esempio.docx schema.png
  2. si riformatta `esempio.docx` in Word
  3. python scripts/tag-relazione-template.py esempio.docx
  4. si rigenera l'esempio e si confronta

Lo schema d'impianto al punto 1 non è facoltativo: senza immagine il paragrafo di §2.3
sparisce del tutto dal documento reso, e non ci sarebbe nulla da sostituire con il tag.

Se un giorno si preferisse curare direttamente il template, basta smettere di eseguire lo
script: da quel momento la sorgente di verità è il `.docx` in `public/templates/`.
"""

import copy
import os
import re
import shutil
import sys
import zipfile

import docx

# Le valvole di un recipiente vanno una per riga dentro la cella: i tag di loop stanno in
# paragrafi propri, che docxtemplater rimuove lasciando solo le righe ripetute. Scritti di
# seguito nello stesso paragrafo ripeterebbero il testo senza mai andare a capo.
VALVOLE_ANNIDATE = ["{#valvole}", "{pos} – n.f. {nFabbrica}", "{/valvole}"]

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# Parti del pacchetto che esistono solo per i commenti di revisione: nel template non
# hanno nulla da fare. `people.xml` è l'anagrafica di chi li ha scritti.
PARTI_COMMENTI = (
    "word/comments.xml",
    "word/commentsExtended.xml",
    "word/commentsIds.xml",
    "word/commentsExtensible.xml",
    "word/people.xml",
)

# La radice del progetto si ricava dalla posizione dello script, non dalla cartella da cui
# lo si lancia: il template sta sempre lì, e chi segue la guida si trova spesso dentro
# `DOCUMENTAZIONE/relazione/`. Il documento sorgente passato a riga di comando resta invece
# relativo alla cartella corrente.
RADICE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

SRC_DEFAULT = os.path.join(RADICE, "DOCUMENTAZIONE", "relazione", "ESEMPIO_nuova_struttura_revFB.docx")
OUT = os.path.join(RADICE, "public", "templates", "relazione-dm329.docx")


# ---------------------------------------------------------------------------
# Primitive
# ---------------------------------------------------------------------------

def atteso(condizione, messaggio):
    if not condizione:
        raise SystemExit("Il documento sorgente non ha la forma attesa: " + messaggio)


def testo(p):
    return p.text.strip()


# ---------------------------------------------------------------------------
# Ricerca degli agganci per testo
#
# Si tengono riferimenti ai paragrafi, non i loro indici: cancellarne uno non invalida
# gli altri, mentre gli indici slitterebbero tutti.
# ---------------------------------------------------------------------------

def paragrafo(P, inizio, cosa):
    """Il primo paragrafo il cui testo inizia con `inizio`. Fallisce se non c'è o è doppio."""
    trovati = [p for p in P if testo(p).startswith(inizio)]
    atteso(trovati, "%s: nessun paragrafo inizia con %r" % (cosa, inizio))
    atteso(
        len(trovati) == 1,
        "%s: %d paragrafi iniziano con %r, l'aggancio è ambiguo"
        % (cosa, len(trovati), inizio),
    )
    return trovati[0]


def precedente(P, p, cosa):
    """
    Il paragrafo che precede `p`, usato come modello di formattazione per i paragrafi di
    tag che verranno inseriti: prende font e stile del corpo del testo, non di un titolo.
    """
    i = P.index(p)
    atteso(i > 0, "%s: non c'è un paragrafo precedente da cui prendere la formattazione" % cosa)
    return P[i - 1]


def seguenti_che_iniziano(P, p, inizio):
    """I paragrafi consecutivi dopo `p` che iniziano con `inizio`, per cancellarli in blocco."""
    fuori = []
    for q in P[P.index(p) + 1:]:
        if not testo(q).startswith(inizio):
            break
        fuori.append(q)
    return fuori


def paragrafi_fra(P, dopo, prima_di, cosa):
    """
    I paragrafi compresi fra due agganci. Delimitare un elenco con i capoversi che lo
    racchiudono, invece che con la forma delle sue voci, lo rende indipendente dal fatto
    che siano trattini scritti a mano o un elenco puntato di Word.
    """
    i, j = P.index(dopo), P.index(prima_di)
    atteso(j > i, "%s: gli agganci sono in ordine inverso" % cosa)
    fra = P[i + 1:j]
    atteso(fra, "%s: nessun paragrafo fra i due agganci" % cosa)
    return fra


def seguenti_con_stile(P, p):
    """
    I paragrafi consecutivi dopo `p` che ne condividono lo stile.

    Serve per gli elenchi le cui voci non hanno un prefisso comune su cui agganciarsi:
    delimitarli con «tutto ciò che segue» cancellerebbe anche un capoverso aggiunto in coda.
    """
    fuori = []
    for q in P[P.index(p) + 1:]:
        if q.style.name != p.style.name:
            break
        fuori.append(q)
    return fuori


def tabella(T, intestazione, cosa):
    """
    La tabella la cui riga di intestazione inizia con le celle indicate.

    L'indice non basta a distinguerle: §4 e §5.2 aprono entrambe con
    «Pos. · Descrizione · Costruttore e modello», e si separano solo dalla quarta colonna.
    """
    def combacia(t):
        celle = [c.text.strip() for c in t.rows[0].cells]
        if len(celle) < len(intestazione):
            return False
        return all(
            attesa is None or celle[j].startswith(attesa)
            for j, attesa in enumerate(intestazione)
        )

    trovate = [t for t in T if t.rows and combacia(t)]
    atteso(trovate, "%s: nessuna tabella con intestazione %r" % (cosa, intestazione))
    atteso(
        len(trovate) == 1,
        "%s: %d tabelle con intestazione %r, l'aggancio è ambiguo"
        % (cosa, len(trovate), intestazione),
    )
    return trovate[0]


def scrivi(p, nuovo, run_modello=None):
    """
    Sostituisce il contenuto del paragrafo conservando la formattazione del primo run.

    Le celle vuote non hanno alcun run: lì il run va creato, e prende le proprietà da un
    run modello della stessa riga — altrimenti il tag erediterebbe il font di base invece
    di quello della tabella.
    """
    if not p.runs:
        r = p.add_run()
        if run_modello is not None and run_modello._r.rPr is not None:
            r._r.insert(0, copy.deepcopy(run_modello._r.rPr))
    p.runs[0].text = nuovo
    for r in list(p.runs[1:]):
        r._r.getparent().remove(r._r)


def elimina(p):
    p._p.getparent().remove(p._p)


def aggiungi_run(p, testo_run, corsivo=False):
    """
    Accoda un run al paragrafo copiando la formattazione dell'ultimo esistente, così il
    testo aggiunto eredita font e corpo del capoverso e non quelli di base del documento.
    """
    r = p.add_run()
    modello = p.runs[0] if p.runs else None
    if modello is not None and modello._r.rPr is not None:
        r._r.insert(0, copy.deepcopy(modello._r.rPr))
    r.text = testo_run
    if corsivo:
        r.italic = True
    return r


def clona_dopo(p, nuovo_testo=None, evidenzia=None):
    """Duplica il paragrafo (formattazione inclusa) e lo inserisce subito dopo."""
    copia = copy.deepcopy(p._p)
    p._p.addnext(copia)
    clone = docx.text.paragraph.Paragraph(copia, p._parent)
    if nuovo_testo is not None:
        scrivi(clone, nuovo_testo)
    if evidenzia is not None:
        for r in clone.runs:
            r.font.highlight_color = evidenzia
    return clone


def paragrafo_tag(modello, tag, dopo):
    """
    Inserisce un paragrafo che contiene *solo* un tag di sezione.

    docxtemplater con `paragraphLoop` elimina i paragrafi che contengono soltanto un tag di
    apertura o chiusura: è il modo di delimitare un blocco senza inquinare il testo reso.
    """
    copia = copy.deepcopy(modello._p)
    dopo._p.addnext(copia)
    p = docx.text.paragraph.Paragraph(copia, modello._parent)
    scrivi(p, tag)
    p.paragraph_format.space_after = 0
    return p


def avvolgi(p, nome_sezione, modello=None):
    """Circonda un paragrafo con {#sezione} … {/sezione}, senza toccarne i run."""
    base = modello or p
    chiusura = paragrafo_tag(base, "{/%s}" % nome_sezione, p)
    apertura = copy.deepcopy(base._p)
    p._p.addprevious(apertura)
    ap = docx.text.paragraph.Paragraph(apertura, p._parent)
    scrivi(ap, "{#%s}" % nome_sezione)
    ap.paragraph_format.space_after = 0
    return ap, chiusura


def avvolgi_blocco(p, t, nome_sezione, modello=None):
    """
    Come `avvolgi`, ma il blocco arriva fino a una tabella: apertura prima del paragrafo,
    chiusura dopo la tabella.

    Serve dove il condizionale non regge una frase ma un pezzo di documento — in §5.3
    l'intestazione «Altre apparecchiature…» e la tabella che la segue stanno o cadono
    insieme: un titolo con sotto una tabella vuota è peggio dell'assenza di entrambi.
    """
    base = modello or p

    apertura = copy.deepcopy(base._p)
    p._p.addprevious(apertura)
    ap = docx.text.paragraph.Paragraph(apertura, p._parent)
    scrivi(ap, "{#%s}" % nome_sezione)
    ap.paragraph_format.space_after = 0

    chiusura = copy.deepcopy(base._p)
    t._tbl.addnext(chiusura)
    ch = docx.text.paragraph.Paragraph(chiusura, p._parent)
    scrivi(ch, "{/%s}" % nome_sezione)
    ch.paragraph_format.space_after = 0

    return ap, ch


def riga_loop(tabella, lista, celle, attesi):
    """
    Riduce una tabella al solo modello di riga: intestazione + una riga di dati, con i tag
    al posto dei valori e il loop che la avvolge.

    `celle` è una lista per colonna: una stringa (cella a un paragrafo) oppure una lista di
    stringhe (cella a più paragrafi, es. costruttore su una riga e modello sull'altra).

    Per i loop annidati dentro una cella (le valvole di un recipiente, i compressori
    connessi a una valvola) i tag vanno in paragrafi propri: `{#valvole}…{/valvole}` scritti
    di seguito nello stesso paragrafo ripeterebbero il testo senza andare a capo.
    """
    atteso(len(tabella.rows) >= 2, "tabella con meno di due righe")
    riga = tabella.rows[1]
    atteso(
        len(riga.cells) == len(celle),
        "attese %d colonne, trovate %d" % (len(celle), len(riga.cells)),
    )
    for j, valore in enumerate(attesi):
        if valore is None:
            continue
        trovato = riga.cells[j].text.strip()
        atteso(trovato == valore, "colonna %d: atteso %r, trovato %r" % (j, valore, trovato))

    for r in list(tabella.rows[2:]):
        tabella._tbl.remove(r._tr)

    # La riga superstite può essere la capofila di una fusione verticale del documento
    # reso. Lasciarla porterebbe `vMerge` in *ogni* riga generata, dove poi `fusioneCelle`
    # ne inietta un'altra: le fusioni le calcola il render, non il template.
    for cella in riga.cells:
        tcPr = cella._tc.tcPr
        if tcPr is None:
            continue
        for vmerge in tcPr.findall(W + "vMerge"):
            tcPr.remove(vmerge)

    # Run di riferimento per le celle vuote: la formattazione del testo di tabella sta sui
    # run, non sullo stile, quindi va presa da una cella che ne ha uno.
    run_modello = next(
        (p.runs[0] for c in riga.cells for p in c.paragraphs if p.runs), None
    )

    for j, contenuto in enumerate(celle):
        cella = riga.cells[j]
        pezzi = contenuto if isinstance(contenuto, list) else [contenuto]
        for k, pezzo in enumerate(pezzi):
            if k < len(cella.paragraphs):
                scrivi(cella.paragraphs[k], pezzo, run_modello)
            else:
                nuovo = cella.add_paragraph()
                scrivi(nuovo, pezzo, run_modello)
        for p in list(cella.paragraphs[len(pezzi):]):
            elimina(p)

    prima = riga.cells[0].paragraphs[0]
    prima.runs[0].text = "{#%s}%s" % (lista, prima.runs[0].text)
    ultima = riga.cells[-1].paragraphs[-1]
    ultima.runs[-1].text = "%s{/%s}" % (ultima.runs[-1].text, lista)


# ---------------------------------------------------------------------------
# Conversione
# ---------------------------------------------------------------------------

def indirizzo_copertina(P, dopo, tag, cosa):
    """
    Sostituisce l'indirizzo di copertina col suo tag, accettando entrambe le forme.

    Nel documento formattato a mano nel 2026 via e località stavano su due paragrafi; da
    quando il motore le unisce in `sedeLegaleCopertina`, un documento reso ne ha uno solo
    con l'a capo dentro (docxtemplater, con `linebreaks: true`, lo rende come <w:br/>).
    Il template vuole in ogni caso un paragrafo solo: se ne trova due, il secondo si toglie.

    Il confronto ignora le maiuscole: il motore rende gli indirizzi tutti maiuscoli, ma il
    documento formattato a mano li porta ancora com'erano stati digitati, e l'una e l'altra
    forma devono restare leggibili.
    """
    i = P.index(dopo)
    atteso(i + 1 < len(P), "%s: manca il paragrafo dell'indirizzo" % cosa)
    riga = P[i + 1]
    atteso(
        testo(riga).upper().startswith("VIA ESEMPIO"),
        "%s: atteso l'indirizzo di esempio, trovato %r" % (cosa, testo(riga)[:40]),
    )
    scrivi(riga, tag)
    # Forma a due paragrafi: la località segue e va assorbita nel tag.
    if i + 2 < len(P) and testo(P[i + 2]).startswith("31013"):
        elimina(P[i + 2])


APERTURA_MOTIVO = "conseguente a:"
CHIUSURA_MOTIVO = ". Vengono verificati"

APERTURA_SPESSIMETRICHE = "verifiche di integrità,"
CHIUSURA_SPESSIMETRICHE = " a verifica spessimetrica"


def segnaposto_fra(p, apertura, chiusura, tag, cosa):
    """
    Sostituisce col tag il testo compreso fra due frasi d'aggancio.

    Ci si aggancia a ciò che circonda il testo variabile, non alla sua forma: nel
    documento formattato a mano lì c'è un segnaposto («[descrivere le motivazioni…]»,
    «XXX»), in uno reso c'è il dato vero, e sono testi diversi. Cercare il segnaposto
    avrebbe chiuso il giro, impedendo di ricavare il template da un documento uscito dal
    template.

    Il capoverso si ricompone in tre run — prefisso, tag, suffisso — prendendo la
    formattazione dai run esistenti. L'evidenziazione non si riporta: dove c'era
    segnalava «qui manca qualcosa da scrivere a mano», e quel qualcosa ora arriva dal
    form come ogni altro dato.
    """
    intero = p.text
    i = intero.find(apertura)
    atteso(i >= 0, "%s: il capoverso non contiene %r" % (cosa, apertura))
    j = intero.find(chiusura, i + len(apertura))
    atteso(j >= 0, "%s: il capoverso non contiene %r dopo il segnaposto" % (cosa, chiusura))

    # Lo spazio fra aggancio e tag si normalizza: il documento reso e quello formattato a
    # mano non sono tenuti a spaziarlo allo stesso modo.
    prefisso = intero[: i + len(apertura)].rstrip() + " "
    suffisso = intero[j:]

    rPr_iniziale = p.runs[0]._r.rPr if p.runs else None
    rPr_finale = p.runs[-1]._r.rPr if p.runs else None
    for r in list(p.runs):
        r._r.getparent().remove(r._r)

    for testo_run, rPr in ((prefisso, rPr_iniziale), (tag, rPr_iniziale), (suffisso, rPr_finale)):
        r = p.add_run()
        if rPr is not None:
            r._r.insert(0, copy.deepcopy(rPr))
        r.text = testo_run
        r.font.highlight_color = None


def paragrafo_immagine(P, dopo):
    """
    Il primo paragrafo con un'immagine dopo `dopo`, scavalcando i paragrafi vuoti.

    Le righe vuote fra il capoverso e lo schema sono spaziatura del redattore e restano
    dove sono: qui servono solo ad attraversarle.
    """
    for q in P[P.index(dopo) + 1:]:
        if "<w:drawing" in q._p.xml:
            return q
        atteso(
            not testo(q),
            "§2.3: fra l'introduzione e il testo %r non c'è l'immagine dello schema. "
            "Genera l'esempio passando un file immagine come secondo argomento: senza "
            "schema quel paragrafo non compare affatto nel documento reso" % testo(q)[:40],
        )
    raise SystemExit(
        "Il documento sorgente non ha la forma attesa: §2.3: nessuna immagine dopo "
        "l'introduzione allo schema d'impianto"
    )


def togli_commenti(d):
    """
    Toglie dal corpo i commenti di revisione.

    Sono appunti sul documento di esempio — «questo paragrafo compare solo se…» — e
    servono a decidere cosa deve fare il template, non a viaggiare dentro ogni relazione
    generata. Restano dove sono utili, nel `.docx` di esempio in `DOCUMENTAZIONE/`.

    Qui si tolgono gli ancoraggi nel testo; le parti del pacchetto che li definiscono le
    lascia fuori `rifinisci_pacchetto`.
    """
    corpo = d.element.body
    quanti = 0

    for tag in ("commentRangeStart", "commentRangeEnd"):
        for el in list(corpo.iter(W + tag)):
            el.getparent().remove(el)

    # Il richiamo numerato è un run a sé: si toglie il run, non il solo riferimento, o
    # resterebbe un run vuoto con lo stile del rimando.
    for rif in list(corpo.iter(W + "commentReference")):
        run = rif.getparent()
        run.getparent().remove(run)
        quanti += 1

    return quanti


def converti(sorgente, destinazione):
    d = docx.Document(sorgente)

    # I commenti si tolgono prima di ogni altra cosa. Il richiamo numerato è un run in
    # coda al capoverso, e le sostituzioni che ricompongono un paragrafo prendono la
    # formattazione dai run esistenti: lasciato lì, il testo del template erediterebbe
    # lo stile «Rimando commento», corpo 10, da un appunto di revisione.
    commenti = togli_commenti(d)

    P = d.paragraphs
    T = d.tables

    # Tutti gli agganci si risolvono prima di modificare alcunché: dopo le cancellazioni
    # `d.paragraphs` non corrisponderebbe più a `P`, mentre i riferimenti restano validi.
    ragione_sociale = paragrafo(P, "ESEMPIO S.P.A.", "copertina")
    sito_produttivo = paragrafo(P, "Sito produttivo in", "copertina")
    premessa = paragrafo(P, "La presente relazione tecnica", "§1 premessa")
    revisione = paragrafo(P, "L’attuale revisione", "§1 capoverso di revisione")
    spessimetriche = paragrafo(P, "Ove previsto", "§1 capoverso spessimetriche")
    intro_sezioni = paragrafo(P, "L’impianto in oggetto è finalizzato", "§2.1 elenco sezioni")
    chiusura_sezioni = paragrafo(P, "L’impianto è protetto contro i rischi", "§2.1 chiusura elenco")
    intro_schema = paragrafo(P, "Lo schema seguente rappresenta", "§2.3 schema d'impianto")
    nocive = paragrafo(P, "Quest’ultima risulta priva", "§3 frase sulle sostanze nocive")
    altre_apparecchiature = paragrafo(
        P, "Altre apparecchiature soggette", "§5.3 altre apparecchiature"
    )
    tabella_altre = tabella(
        T,
        ["Pos.", "Valvole di sicurezza", "Manometro"],
        "§5.3 protezioni delle altre apparecchiature",
    )
    tubazioni = paragrafo(P, "Tutte le tubazioni", "§5.4 nota sulle tubazioni")
    intro_riqualificazione = paragrafo(
        P, "Le attrezzature rientranti nel campo", "§7 introduzione"
    )
    spessimetriche_svolte = paragrafo(
        P, "Come già evidenziato nella tabella", "§7.2 capoverso spessimetriche"
    )
    primo_allegato = paragrafo(P, "Attestazioni", "§8 elenco allegati")

    # --- Copertina ---------------------------------------------------------
    scrivi(ragione_sociale, "{premessa.ragioneSociale}")
    indirizzo_copertina(
        P, ragione_sociale, "{premessa.sedeLegaleCopertina}", "copertina · sede legale"
    )
    indirizzo_copertina(
        P, sito_produttivo, "{premessa.sitoProduttivoCopertina}", "copertina · sito produttivo"
    )

    # --- §1 Premessa -------------------------------------------------------
    # La denominazione della sala va in corsivo, quindi non può stare nello stesso run del
    # resto della frase: il capoverso si compone in quattro run, uno solo dei quali corsivo.
    scrivi(
        premessa,
        "La presente relazione tecnica si riferisce all’impianto a pressione installato "
        "presso il sito produttivo della ditta {premessa.ragioneSociale}, con sede sociale "
        "in {premessa.sedeLegale}, esercente attività di {premessa.descrizioneAttivita}, "
        "{premessa.ubicazione}",
    )
    aggiungi_run(premessa, "{#premessa.haDenominazioneSala} ed individuato come ")
    aggiungi_run(premessa, "{premessa.denominazioneSala}", corsivo=True)
    aggiungi_run(premessa, "{/premessa.haDenominazioneSala}.")

    # Dei due capoversi condizionali si tocca un run solo: quello del motivo della
    # revisione, che il redattore scrive ora nel form invece che a mano in Word. Il resto
    # della frase resta com'è, perché riscriverla perderebbe la formattazione del
    # capoverso. Via anche il giallo: segnalava «qui manca qualcosa da scrivere», e ora
    # quel qualcosa arriva dal form come tutti gli altri dati.
    segnaposto_fra(
        revisione,
        APERTURA_MOTIVO,
        CHIUSURA_MOTIVO,
        "{premessa.motivoRevisione}",
        "§1 capoverso di revisione",
    )
    modello_premessa = precedente(P, revisione, "§1 capoverso di revisione")
    avvolgi(revisione, "premessa.haRevisione", modello=modello_premessa)
    avvolgi(spessimetriche, "premessa.haSpessimetrica", modello=modello_premessa)

    # --- §2.1 Sezioni dell'impianto ---------------------------------------
    # L'elenco è delimitato dai due capoversi che lo racchiudono, non dalla forma delle
    # voci: possono essere paragrafi con un trattino scritto a mano oppure un elenco
    # puntato di Word, e la scelta è del redattore.
    voci_sezioni = paragrafi_fra(P, intro_sezioni, chiusura_sezioni, "§2.1 elenco sezioni")
    prefisso = "–\t" if testo(voci_sezioni[0]).startswith("–") else ""
    scrivi(voci_sezioni[0], prefisso + "{voce}")
    avvolgi(voci_sezioni[0], "descrizioneGenerale.sezioni", modello=intro_sezioni)
    for voce in voci_sezioni[1:]:
        elimina(voce)

    # --- §2.2 Condizioni di installazione ----------------------------------
    riga_loop(
        tabella(T, ["Requisito", "Esito"], "§2.2 condizioni di installazione"),
        "condizioniInstallazione",
        ["{requisito}", "{esito}"],
        ["Ubicazione impianto", None],
    )

    # --- §2.3 Schema d'impianto -------------------------------------------
    # L'immagine si cerca scavalcando i paragrafi vuoti: spaziare l'introduzione dallo
    # schema è formattazione, e pretenderla nel paragrafo immediatamente successivo
    # rimetterebbe un aggancio posizionale in mezzo a quelli per testo. Un paragrafo con
    # del testo, invece, chiude la ricerca: oltre comincia un'altra sezione.
    scrivi(paragrafo_immagine(P, intro_schema), "{%schemaImpianto}")

    # --- §3 Fluidi ---------------------------------------------------------
    riga_loop(
        tabella(T, ["Circuito", "Fluido", "Gruppo", "Provenienza"], "§3 fluidi di processo"),
        "fluidi.righe",
        ["{circuito}", "{fluido}", "{gruppo}", "{provenienza}"],
        ["Aria compressa", "Aria ambiente", None, None],
    )

    # La frase esiste in due varianti: evidenziata quando l'aria aspirata è dichiarata non
    # pulita, piana altrimenti. Il documento reso ne contiene una sola, e qui si ricreano
    # entrambe a partire da quella.
    #
    # L'evidenziatore lo mette lo script, senza guardare come sia formattata la frase nel
    # documento sorgente: è una marcatura che il sistema deve al lettore — «qui c'è una
    # valutazione da fare» — non una scelta tipografica del redattore. Prima si pretendeva
    # di trovarla già gialla, e toglierla in Word bloccava la conversione.
    modello_fluidi = precedente(P, nocive, "§3 frase sulle sostanze nocive")
    piana = clona_dopo(nocive)
    for r in piana.runs:
        r.font.highlight_color = None
    for r in nocive.runs:
        r.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
    avvolgi(nocive, "fluidi.evidenziaNocive", modello=modello_fluidi)
    avvolgi(piana, "fluidi.evidenziaNocive", modello=modello_fluidi)
    # La sezione inversa non si scrive con {#…}: si corregge il tag di apertura.
    apertura_piana = piana._p.getprevious()
    docx.text.paragraph.Paragraph(apertura_piana, piana._parent).runs[0].text = (
        "{^fluidi.evidenziaNocive}"
    )

    # --- §4 Caratterizzazione ---------------------------------------------
    riga_loop(
        tabella(
            T,
            ["Pos.", "Descrizione", "Costruttore e modello", "Capacit"],
            "§4 caratterizzazione",
        ),
        "caratteristiche",
        ["{pos}", "{descrizione}", ["{costruttore}", "{modello}"], "{capacita}",
         "{pressione}", "{temperatura}", "{categoria}", "{anno}", "{nFabbrica}"],
        ["C1", "Compressore", None, "8000", None, None, None, "2025", None],
    )

    # --- §5.2 Esiti --------------------------------------------------------
    riga_loop(
        tabella(T, ["Pos.", "Descrizione", "Costruttore e modello", "V"], "§5.2 esiti"),
        "esiti",
        ["{pos}", "{apparecchiatura}", ["{costruttore}", "{modello}"], "{volume}", "{ps}",
         "{psPerV}", "{categoria}", "{adempimento}", "{statoInail}", "{verificaIntegritaMark}"],
        ["C1", "Compressore", None, None, None, None, None, "Escluso", "Nuova richiesta", None],
    )

    # --- §5.3 Protezioni ---------------------------------------------------
    riga_loop(
        tabella(
            T,
            ["Pos.", "Valvole di sicurezza", "Scarico condensa"],
            "§5.3 protezioni dei serbatoi",
        ),
        "protezioni.serbatoi",
        ["{pos}", VALVOLE_ANNIDATE, "{scaricoCondensa}",
         "{finituraInterna}", "{ancoraggio}", "{manometro}"],
        ["S1", None, "Automatico", "Zincato", "Sì", None],
    )
    riga_loop(
        tabella_altre,
        "protezioni.altre",
        ["{pos}", VALVOLE_ANNIDATE, "{manometro}"],
        ["C1.1", None, "a bordo macchina"],
    )
    # Impianti senza disoleatori, scambiatori né recipienti filtro esistono: lì
    # l'intestazione e la tabella spariscono insieme. Un booleano e non `{#protezioni.altre}`,
    # che ripeterebbe l'intestazione una volta per riga.
    avvolgi_blocco(altre_apparecchiature, tabella_altre, "protezioni.haAltre")

    # --- §5.4 Tubazioni ----------------------------------------------------
    modello_tubazioni = precedente(P, tubazioni, "§5.4 nota sulle tubazioni")
    oltre = clona_dopo(
        tubazioni,
        "Le tubazioni dell’impianto presentano DN massimo pari a {tubazioni.dnMassimo} mm, "
        "superiore alla soglia di 80 mm dell’art. 3 comma bb): rientrano pertanto nel campo "
        "di applicazione del D.M. 329/2004 e sono soggette ai relativi obblighi di denuncia.",
        evidenzia=docx.enum.text.WD_COLOR_INDEX.YELLOW,
    )
    avvolgi(tubazioni, "tubazioni.escluse", modello=modello_tubazioni)
    avvolgi(oltre, "tubazioni.escluse", modello=modello_tubazioni)
    apertura_oltre = oltre._p.getprevious()
    docx.text.paragraph.Paragraph(apertura_oltre, oltre._parent).runs[0].text = (
        "{^tubazioni.escluse}"
    )

    # --- §6 Valvole di sicurezza ------------------------------------------
    # Tre righe per apparecchiatura connessa: posizione e descrizione, poi costruttore e
    # modello a capo. Le celle sono strette e il nome del costruttore è lungo.
    connesse = [
        "{#connesse}",
        "Pos. {pos} – {descrizione} –",
        "{costruttore}",
        "{modello}",
        "{/connesse}",
    ]
    riga_loop(
        tabella(
            T,
            ["Pos. valvola", "N. fabbrica", None, "Portata massima"],
            "§6.1 portata delle valvole",
        ),
        "valvole.portata",
        ["{posValvola}", "{nFabbricaValvola}", connesse, "{portataMaxTesto}",
         "{portataScaricata}", "{adeguatoMark}"],
        ["C1.2", "1001913158", None, "8000", "21500", None],
    )
    riga_loop(
        tabella(
            T,
            ["Pos. valvola", "N. fabbrica", None, "PS recipiente"],
            "§6.2 pressione di taratura",
        ),
        "valvole.pressione",
        ["{posValvola}", "{nFabbricaValvola}", connesse, "{psRecipiente}",
         "{pressioneTaratura}", "{adeguatoMark}"],
        ["C1.2", "1001913158", None, "11", "10", None],
    )

    # --- §7.2 Riqualificazione --------------------------------------------
    riga_loop(
        # Non «Attrezzature contenenti…»: quella è la tabella delle frequenze di §7.1,
        # che ha intestazioni quasi identiche ma nessuna riga da generare.
        tabella(
            T,
            ["Pos.", "Apparecchiatura", "Cat.", "Verifica di funzionamento"],
            "§7.2 scadenze di riqualificazione",
        ),
        "riqualificazione",
        ["{pos}", "{apparecchiatura}", "{categoria}", "{verificaFunzionamento}",
         "{verificaIntegrita}"],
        ["C1.1", "Serbatoio disoleatore", "III", None, None],
    )

    # Il capoverso che chiude §7.2 nomina le apparecchiature già verificate: senza
    # nessuna da nominare non viene stampato affatto. La clausola arriva accordata al
    # numero dal motore — «l'apparecchiatura S1 è stata sottoposta» oppure «le
    # apparecchiature C2.1 e S1 sono state sottoposte» — così il template resta muto.
    segnaposto_fra(
        spessimetriche_svolte,
        APERTURA_SPESSIMETRICHE,
        CHIUSURA_SPESSIMETRICHE,
        "{spessimetriche.clausola}",
        "§7.2 capoverso spessimetriche",
    )
    avvolgi(
        spessimetriche_svolte, "spessimetriche.presenti", modello=intro_riqualificazione
    )

    # --- §8 Allegati -------------------------------------------------------
    # Le altre voci si cancellano una a una: sono capoversi di elenco, senza un prefisso
    # comune su cui agganciarsi, e stanno fra la prima voce e la fine del documento.
    altre_voci = seguenti_con_stile(P, primo_allegato)
    # Modello dal corpo del testo e non dalla voce di elenco: i tag erediterebbero il
    # rientro e il punto elenco.
    scrivi(primo_allegato, "{voce}")
    avvolgi(primo_allegato, "allegati", modello=intro_riqualificazione)
    for voce in altre_voci:
        elimina(voce)

    tabella_revisioni(d)

    d.save(destinazione)
    return commenti


def tabella_revisioni(d):
    """
    Tabella delle revisioni, nel piè di pagina della copertina: l'ultima riga riporta la
    data di emissione scelta nel form, il numero di revisione desunto dal codice pratica e,
    alla prima emissione, la nota corrispondente. Le righe superiori restano vuote: le
    compila il tecnico a ogni revisione.
    """
    trovate = [
        t
        for s in d.sections
        for f in (s.footer, s.first_page_footer, s.even_page_footer)
        for t in f.tables
        if [c.text.strip() for c in t.rows[0].cells][:3] == ["DATA", "REV", "OGGETTO"]
    ]
    atteso(trovate, "tabella delle revisioni: non trovata nei piè di pagina")
    atteso(
        len(trovate) == 1,
        "tabella delle revisioni: trovate %d tabelle, l'aggancio è ambiguo" % len(trovate),
    )

    ultima = trovate[0].rows[-1]
    atteso(len(ultima.cells) >= 3, "tabella delle revisioni: attese almeno 3 colonne")
    modello = next((p.runs[0] for c in ultima.cells for p in c.paragraphs if p.runs), None)
    scrivi(ultima.cells[0].paragraphs[0], "{premessa.dataEmissione}", modello)
    scrivi(ultima.cells[1].paragraphs[0], "{premessa.numeroRevisione}", modello)
    scrivi(ultima.cells[2].paragraphs[0], "{premessa.notaRevisione}", modello)


def senza_relazioni(rels, target):
    """Toglie dal file delle relazioni quelle che puntano a una parte rimossa."""
    return re.sub(
        r'<Relationship[^>]*Target="%s"[^>]*/>' % re.escape(target), "", rels
    )


def rifinisci_pacchetto(percorso):
    """
    Tre ritocchi che vivono solo nel pacchetto zip, non nel modello di python-docx:

    - «Pag. 2di 9»: nel piè di pagina manca lo spazio prima di «di». Refuso di battitura,
      si corregge qui perché il template è ormai la sorgente di verità.
    - L'immagine dello schema del documento reso resta orfana una volta sostituita dal tag:
      va tolta dal pacchetto insieme alla sua relazione, altrimenti verrebbe trascinata in
      ogni relazione generata.
    - Le parti dei commenti di revisione: `togli_commenti` ne ha tolto gli ancoraggi dal
      testo, qui se ne vanno anche le definizioni. Un Override che punta a una parte
      assente rende il documento illeggibile, quindi vanno ripuliti anche i tipi di
      contenuto e le relazioni.
    """
    temporaneo = percorso + ".tmp"
    esiti = {"spazio": False, "immagine": False, "commenti": False}

    with zipfile.ZipFile(percorso) as src:
        nomi = set(src.namelist())
        media = sorted(n for n in nomi if n.startswith("word/media/"))
        commenti = sorted(n for n in PARTI_COMMENTI if n in nomi)
        documento = src.read("word/document.xml").decode("utf-8")
        atteso(
            'r:embed="' not in documento,
            "il template contiene ancora un riferimento a un'immagine",
        )

        rels = src.read("word/_rels/document.xml.rels").decode("utf-8")
        for parte in media + commenti:
            rels = senza_relazioni(rels, parte[len("word/"):])

        tipi = src.read("[Content_Types].xml").decode("utf-8")
        for parte in commenti:
            tipi = re.sub(
                r'<Override[^>]*PartName="/%s"[^>]*/>' % re.escape(parte), "", tipi
            )

        with zipfile.ZipFile(temporaneo, "w", zipfile.ZIP_DEFLATED) as dst:
            for info in src.infolist():
                if info.filename in media:
                    esiti["immagine"] = True
                    continue
                if info.filename in commenti:
                    esiti["commenti"] = True
                    continue

                dati = src.read(info.filename)
                if info.filename.startswith("word/footer"):
                    xml = dati.decode("utf-8")
                    if '<w:t xml:space="preserve">di </w:t>' in xml:
                        xml = xml.replace(
                            '<w:t xml:space="preserve">di </w:t>',
                            '<w:t xml:space="preserve"> di </w:t>',
                        )
                        dati = xml.encode("utf-8")
                        esiti["spazio"] = True
                elif info.filename == "word/_rels/document.xml.rels":
                    dati = rels.encode("utf-8")
                elif info.filename == "[Content_Types].xml":
                    dati = tipi.encode("utf-8")

                dst.writestr(info, dati)

    shutil.move(temporaneo, percorso)
    return esiti


if __name__ == "__main__":
    sorgente = sys.argv[1] if len(sys.argv) > 1 else SRC_DEFAULT
    commenti = converti(sorgente, OUT)
    esiti = rifinisci_pacchetto(OUT)
    print("Template scritto in %s (%d byte)" % (OUT, os.path.getsize(OUT)))
    print("  spazio nel pie di pagina: %s" % ("corretto" if esiti["spazio"] else "non trovato"))
    print("  immagine di esempio: %s" % ("rimossa" if esiti["immagine"] else "assente"))
    print(
        "  commenti di revisione: %s"
        % ("%d rimossi" % commenti if esiti["commenti"] else "assenti")
    )
